#!/usr/bin/env python3
"""
Export figures and supporting tables with PLOS ONE naming conventions.
"""

from pathlib import Path
import shutil

import pandas as pd
import matplotlib.pyplot as plt
from PIL import Image
import numpy as np
from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[2]
FIG_DIR = REPO_ROOT / "outputs" / "figures"
TABLE_DIR = REPO_ROOT / "outputs" / "tables"
OUT_DIR = REPO_ROOT / "outputs" / "plos_submission"


def csv_to_docx(csv_path, docx_path, title):
    df = pd.read_csv(csv_path)
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])
    doc.save(docx_path)

def copy_or_convert_tiff(stem_path: Path, dest_path: Path) -> bool:
    tiff_path = stem_path.with_suffix(".tiff")
    png_path = stem_path.with_suffix(".png")
    if tiff_path.exists():
        shutil.copyfile(tiff_path, dest_path)
        return True
    if png_path.exists():
        img = plt.imread(png_path)
        if img.dtype != np.uint8:
            img = (img * 255).clip(0, 255).astype(np.uint8)
        Image.fromarray(img).save(dest_path, format="TIFF")
        return True
    return False


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    # Main figures
    for i in range(1, 8):
        src = FIG_DIR / f"figure{i}_"
        matches = list(FIG_DIR.glob(f"figure{i}_*_final.tiff"))
        if matches:
            dest = OUT_DIR / f"Fig{i}.tif"
            shutil.copyfile(matches[0], dest)
            print(f"Copied {dest}")
        else:
            print(f"Missing TIFF for figure {i}")

    # Supporting figure
    s1_png = FIG_DIR / "sample_size_analysis.png"
    if s1_png.exists():
        img = plt.imread(s1_png)
        if img.dtype != np.uint8:
            img = (img * 255).clip(0, 255).astype(np.uint8)
        Image.fromarray(img).save(OUT_DIR / "S1_Fig.tif", format="TIFF")
        print("Created S1_Fig.tif")

    # Workflow flowchart as supporting figure.
    flowchart_candidates = [
        FIG_DIR / "flowchart_pipeline_formatted",
        FIG_DIR / "flowchart_pipeline_final",
    ]
    flowchart_ok = False
    for stem in flowchart_candidates:
        if copy_or_convert_tiff(stem, OUT_DIR / "S2_Fig.tif"):
            print(f"Created S2_Fig.tif from {stem.name}")
            flowchart_ok = True
            break
    if not flowchart_ok:
        print("Missing flowchart pipeline figure (formatted/final)")

    # ERA5 downsampling curve as supporting figure.
    era5_stem = FIG_DIR / "era5_downsample_curve"
    if copy_or_convert_tiff(era5_stem, OUT_DIR / "S4_Fig.tif"):
        print("Created S4_Fig.tif")
    else:
        print("Missing era5_downsample_curve.(tiff/png)")

    # Hydrographic target distribution as supporting figure.
    hydro_png = FIG_DIR / "hydrographic_target_distribution.png"
    if hydro_png.exists():
        img = plt.imread(hydro_png)
        if img.dtype != np.uint8:
            img = (img * 255).clip(0, 255).astype(np.uint8)
        Image.fromarray(img).save(OUT_DIR / "S3_Fig.tif", format="TIFF")
        print("Created S3_Fig.tif")
    else:
        print("Missing hydrographic_target_distribution.png")


    # Supporting tables
    table_map = {
        "S1_Table.docx": "final_table2_model_performance.csv",
        "S2_Table.docx": "complete_sanity_check_results.csv",
        "S3_Table.docx": "small_sample_analysis.csv",
        "S4_Table.docx": "era5_downsample_ablation.csv",
        "S5_Table.docx": "alternative_metrics_detailed.csv",
        "S6_Table.docx": "dataset_split_summary.csv",
        "S7_Table.docx": "dataset_geography_summary.csv",
        "S8_Table.docx": "transfer_pilot.csv",
    }
    for out_name, src_name in table_map.items():
        src_path = TABLE_DIR / src_name
        if src_path.exists():
            csv_to_docx(src_path, OUT_DIR / out_name, out_name.replace("_", " ").replace(".docx", ""))
            print(f"Created {out_name}")
        else:
            print(f"Missing table: {src_name}")


if __name__ == "__main__":
    main()
